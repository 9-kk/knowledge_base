import os
import uuid
from datetime import datetime
import requests

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import aiofiles
from docx import Document
from pydantic import BaseModel
from typing import List
from datetime import date
from minio.error import S3Error
import tempfile
import logging
import uvicorn

from utils.docx_utils import insert_image

# 环境变量配置
from dotenv import load_dotenv
load_dotenv()  # 加载.env文件

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from minio import Minio
# 从环境变量获取配置
MINIO_PORT = os.getenv("MINIO_PORT", "zoomtech2008.cn:9001")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY", "zhongen-oa")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY", "nF!P1rPblFJr#NaZ")
BUCKET_NAME = os.getenv("MINIO_BUCKET", "zhongen-oa-dev")
CALLBACK_URL = os.getenv("CALLBACK_URL", "http://192.168.15.40:8080/system/projectSurveyDocumentManage/report/recall")

minio_client = Minio(
    endpoint=MINIO_PORT,
    access_key=MINIO_ACCESS_KEY,
    secret_key=MINIO_SECRET_KEY,
    secure=True  # 设为True如果使用HTTPS
)


# # ----------配置minio文件存储--------------
# # MINIO在本机上的路径
# MINIO_PORT = "zoomtech2008.cn:9001"
# PROTOCOL = "https://"
# minio_client = Minio(endpoint=MINIO_PORT,
#                      access_key="zhongen-oa",
#                      secret_key="nF!P1rPblFJr#NaZ",
#                      secure=True    # 设为True如果使用HTTPS
#                      )
# BUCKET_NAME = "zhongen-oa-dev"


app = FastAPI()

# 允许跨域请求
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ImageRequest(BaseModel):
    image_names: List[str]
    image_paths: List[str]
    status: int


# 回调后端函数
def callback(callback_data):
    try:
        response = requests.post(CALLBACK_URL, json=callback_data)
        response.raise_for_status()  # 如果响应状态码不是 200，会抛出异常
        print("回调成功:", response.json())
    except requests.exceptions.RequestException as e:
        print(f"回调失败: {e}")


@app.post("/generate-doc")
async def generate_document(request: ImageRequest):

    """生成包含图片的Word文档并上传到MinIO"""
    doc = Document()
    '''
    后续添加文字部分，将页面设置为纵向
    '''
    # section = doc.sections[0]  # 获取第一个节（默认节）
    # section.orientation = 0  # 0 表示纵向
    # section.page_width = Inches(8.27)  # A4 纸纵向宽度
    # section.page_height = Inches(11.69)  # A4 纸纵向高度
    #
    # # 添加一些文字内容
    # doc.add_paragraph("这是纵向页面的文字部分。")
    # doc.add_paragraph("这部分内容将显示在纵向页面上。")

    # 获取当天日期
    today = date.today()
    year = str(today.year)
    month = str(today.month).zfill(2)  # 如果小于10，自动补零
    day = str(today.day).zfill(2)

    # # 临时存储文件夹
    # temp_dir = os.path.join('tmp', year + '_' + month + '_' + day)
    # # 判断文件夹是否存在
    # if not os.path.exists(temp_dir):
    #     # 如果不存在，则新建文件夹
    #     os.makedirs(temp_dir)
    with tempfile.TemporaryDirectory(prefix=f"{year}_{month}_{day}_") as temp_dir:

        temp_files = []  # 用于跟踪创建的临时文件
        try:
            # 下载图片并插入文档
            for img_path in request.image_paths:
                local_path = os.path.join(temp_dir, os.path.basename(img_path))

                # 从MinIO下载图片到临时文件存储空间
                try:
                    minio_client.fget_object(BUCKET_NAME, img_path, local_path)
                    temp_files.append(local_path)
                except S3Error as e:
                    msg = f"Image not found in MinIO: {img_path}. Error: {e.message}"
                    callback_data = {
                        "code": 404,
                        "msg": msg,
                        "data": {
                            "id": request.status,
                            "objectKey": "",
                        }
                    }
                    callback(callback_data)
                    raise HTTPException(
                        status_code=404,
                        detail=msg
                    )

            # 在文档中插入图片
            insert_image(doc, temp_files, request.image_names)
            # # 插入Word文档
            # doc.add_picture(temp_file, width=Inches(6))
            # doc.add_page_break()

            # 保存Word文档到临时文件
            doc_name = f"{uuid.uuid4()}.docx"
            local_doc_path = os.path.join(temp_dir, doc_name)
            doc.save(local_doc_path)

            # 上传到MinIO
            docx_name = f"/common/{year}/{month}/{day}/{doc_name}"
            minio_client.fput_object(
                BUCKET_NAME,
                docx_name,
                local_doc_path,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # 执行成功回调函数
            callback_data = {
                "code": 200,
                "msg": '操作成功',
                "data": {
                    "id": request.status,
                    "objectKey": f"{docx_name}",
                }
            }
            callback(callback_data)

            # 本地调试
            return {
                "message": "Document generated successfully",
                "minio_path": f"{docx_name}",
                "status": request.status
                # "presigned_url": minio_client.presigned_get_object(BUCKET_NAME, docx_name)
            }

        except Exception as e:
            callback_data = {
                "code": 500,
                "msg": str(e),
                "data": {
                    "id": request.status,
                    "objectKey": "",
                }
            }
            callback(callback_data)
            raise HTTPException(status_code=500, detail=str(e))

    # finally:
    #     # 清理临时文件（实际生产环境建议使用tempfile）
    #     for file in os.listdir(temp_dir):
    #         os.remove(os.path.join(temp_dir, file))
    #     os.rmdir(temp_dir)


# ------------------------------------------------------------------------------
# ------------------------国家2000->经纬度------------------------------
# ------------------------------------------------------------------------------
from utils.gzcj_2_gcj02 import gj2000_to_wgs84, decimal_to_dms
from utils.coordTransform_utils import wgs84_to_gcj02
import json
import numpy as np


class GpsRequest(BaseModel):
    project_list: List[List[float]]


@app.post("/GZCJ_2_GCJ02")
async def GZCJ_2_GCJ02(request: GpsRequest):
    try:
        # 转成numpy的array格式，进行并行运算
        # {‘project_list’:[[X1,Y1] , [X2,Y2]]}
        project_array = np.array(request.project_list)
        x = project_array[:, 0]
        y = project_array[:, 1]
        # print(x, y)
        # 广州城建转wgs84
        x, y = gj2000_to_wgs84(x, y)
        new_project_array = np.stack((x, y), axis=-1)
        tmp_project_list = new_project_array.tolist()

        # wgs84转高德地图（因为运算比较复杂，这里就没办法并行计算啦，但是也很快）
        new_project_list = []
        for each_coord in tmp_project_list:
            x, y = each_coord
            # 获取度分秒格式经纬度
            x, y = wgs84_to_gcj02(x, y)
            # print(decimal_to_dms(x), decimal_to_dms(y))
            new_project_list.append([x, y])

        return JSONResponse({'code': 0, 'errmsg': '坐标转换成功', 'project_list': new_project_list})

    except Exception as e:
        logger.error(e)
        return JSONResponse({'code': 400, 'errmsg': '坐标转换失败', '失败原因': e})

# @app.on_event("startup")
# async def init_minio_buckets():
#     """确保MinIO桶存在"""
#     for bucket in BUCKET_NAME:
#         if not minio_client.bucket_exists(bucket):
#             minio_client.make_bucket(bucket)

# uvicorn QASystem.project_servey:app --reload --host 0.0.0.0 --port 8000
# python -m uvicorn project_servey:app --reload --host 0.0.0.0 --port 8000


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

